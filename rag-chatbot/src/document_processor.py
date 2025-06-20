import os
import logging
from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import tiktoken
from dataclasses import dataclass

# Document processing imports
import pypdf
from docx import Document
import csv

from config.settings import settings
from src.utils import validate_file, get_file_hash

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """Represents a chunk of text from a document."""
    text: str
    metadata: Dict[str, Any]
    chunk_id: str
    
    
@dataclass
class ProcessedDocument:
    """Represents a fully processed document."""
    file_path: str
    file_hash: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    total_chunks: int
    total_tokens: int


class BaseDocumentProcessor(ABC):
    """Abstract base class for document processors."""
    
    def __init__(self):
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
        self.encoding = tiktoken.get_encoding("cl100k_base")
        
    @abstractmethod
    def extract_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from document. Returns (text, metadata)."""
        pass
    
    def count_tokens(self, text: str) -> int:
        """Count tokens in text using tiktoken."""
        return len(self.encoding.encode(text))
    
    def chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """
        Split text into chunks with overlap.
        
        Args:
            text: The text to chunk
            metadata: Document metadata to include with each chunk
            
        Returns:
            List of DocumentChunk objects
        """
        chunks = []
        
        # Clean text
        text = text.strip()
        if not text:
            return chunks
        
        # Tokenize the entire text
        tokens = self.encoding.encode(text)
        
        # Calculate chunk positions
        chunk_starts = list(range(0, len(tokens), self.chunk_size - self.chunk_overlap))
        
        for i, start in enumerate(chunk_starts):
            end = min(start + self.chunk_size, len(tokens))
            chunk_tokens = tokens[start:end]
            
            # Decode back to text
            chunk_text = self.encoding.decode(chunk_tokens)
            
            # Create chunk metadata
            chunk_metadata = metadata.copy()
            chunk_metadata.update({
                'chunk_index': i,
                'chunk_start_token': start,
                'chunk_end_token': end,
                'chunk_token_count': len(chunk_tokens),
                'total_chunks': len(chunk_starts)
            })
            
            # Create chunk ID
            chunk_id = f"{metadata.get('file_hash', 'unknown')}_{i}"
            
            chunks.append(DocumentChunk(
                text=chunk_text,
                metadata=chunk_metadata,
                chunk_id=chunk_id
            ))
        
        return chunks
    
    def process(self, file_path: str) -> ProcessedDocument:
        """Process a document and return structured data."""
        # Validate file
        is_valid, error = validate_file(file_path)
        if not is_valid:
            raise ValueError(f"Invalid file: {error}")
        
        # Get file hash
        file_hash = get_file_hash(file_path)
        
        # Extract text and metadata
        text, doc_metadata = self.extract_text(file_path)
        
        # Add file metadata
        doc_metadata.update({
            'file_path': file_path,
            'file_name': os.path.basename(file_path),
            'file_hash': file_hash,
            'file_size': os.path.getsize(file_path)
        })
        
        # Chunk the text
        chunks = self.chunk_text(text, doc_metadata)
        
        # Calculate total tokens
        total_tokens = sum(chunk.metadata['chunk_token_count'] for chunk in chunks)
        
        return ProcessedDocument(
            file_path=file_path,
            file_hash=file_hash,
            chunks=chunks,
            metadata=doc_metadata,
            total_chunks=len(chunks),
            total_tokens=total_tokens
        )


class PDFProcessor(BaseDocumentProcessor):
    """Processor for PDF documents."""
    
    def extract_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file."""
        text_parts = []
        metadata = {'file_type': 'pdf', 'pages': []}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                metadata['total_pages'] = num_pages
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                        metadata['pages'].append({
                            'page_number': page_num + 1,
                            'has_text': True
                        })
                    else:
                        metadata['pages'].append({
                            'page_number': page_num + 1,
                            'has_text': False
                        })
                
                # Get document info if available
                if pdf_reader.metadata:
                    metadata['pdf_metadata'] = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', '')
                    }
        
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
        
        return '\n\n'.join(text_parts), metadata


class TextProcessor(BaseDocumentProcessor):
    """Processor for plain text files."""
    
    def extract_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text file."""
        metadata = {'file_type': 'text'}
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            # Count lines
            lines = text.split('\n')
            metadata['total_lines'] = len(lines)
            metadata['non_empty_lines'] = len([line for line in lines if line.strip()])
            
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            raise ValueError(f"Failed to read text file: {str(e)}")
        
        return text, metadata


class MarkdownProcessor(TextProcessor):
    """Processor for Markdown files."""
    
    def extract_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from Markdown file."""
        text, metadata = super().extract_text(file_path)
        metadata['file_type'] = 'markdown'
        
        # Count headers
        headers = [line for line in text.split('\n') if line.strip().startswith('#')]
        metadata['header_count'] = len(headers)
        
        return text, metadata


class DocxProcessor(BaseDocumentProcessor):
    """Processor for DOCX documents."""
    
    def extract_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file."""
        metadata = {'file_type': 'docx'}
        text_parts = []
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            text_parts.extend(paragraphs)
            metadata['paragraph_count'] = len(paragraphs)
            
            # Extract tables
            table_count = 0
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        table_text.append(' | '.join(row_text))
                
                if table_text:
                    text_parts.append('\n'.join(table_text))
                    table_count += 1
            
            metadata['table_count'] = table_count
            
            # Get document properties
            core_props = doc.core_properties
            if core_props:
                metadata['docx_metadata'] = {
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'created': str(core_props.created) if core_props.created else ''
                }
                
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
        
        return '\n\n'.join(text_parts), metadata


class CSVProcessor(BaseDocumentProcessor):
    """Processor for CSV files."""
    
    def extract_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from CSV file."""
        metadata = {'file_type': 'csv'}
        text_parts = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                csv_reader = csv.reader(file)
                rows = list(csv_reader)
                
                metadata['total_rows'] = len(rows)
                metadata['total_columns'] = len(rows[0]) if rows else 0
                
                # Convert to text format
                for row in rows:
                    if any(cell.strip() for cell in row):
                        text_parts.append(' | '.join(row))
                
        except Exception as e:
            logger.error(f"Error reading CSV file {file_path}: {e}")
            raise ValueError(f"Failed to read CSV file: {str(e)}")
        
        return '\n'.join(text_parts), metadata


class DocumentProcessorFactory:
    """Factory for creating document processors."""
    
    _processors = {
        '.pdf': PDFProcessor,
        '.txt': TextProcessor,
        '.md': MarkdownProcessor,
        '.docx': DocxProcessor,
        '.csv': CSVProcessor
    }
    
    @classmethod
    def get_processor(cls, file_path: str) -> BaseDocumentProcessor:
        """Get appropriate processor for file type."""
        file_ext = Path(file_path).suffix.lower()
        
        processor_class = cls._processors.get(file_ext)
        if not processor_class:
            raise ValueError(f"No processor available for file type: {file_ext}")
        
        return processor_class()
    
    @classmethod
    def process_document(cls, file_path: str) -> ProcessedDocument:
        """Process a document using the appropriate processor."""
        processor = cls.get_processor(file_path)
        return processor.process(file_path)